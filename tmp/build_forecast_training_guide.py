from __future__ import annotations

from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "QuietMel_Crowd_Forecast_Model_Training_Guide_CN.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "0B6F6D"
PALE_BLUE = "E8EEF5"
PALE_TEAL = "E7F3F1"
PALE_YELLOW = "FFF4CE"
PALE_RED = "FDE9E7"
LIGHT_GREY = "F5F6F7"
MID_GREY = "D9DEE3"
DARK = "20262E"
MUTED = "5F6B76"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def format_table(table, widths: list[int], header=True, font_size=8.5) -> None:
    set_table_fixed(table, widths)
    for r_idx, row in enumerate(table.rows):
        keep_row_together(row)
        for cell in row.cells:
            if header and r_idx == 0:
                set_cell_shading(cell, PALE_BLUE)
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.08
                for run in para.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = RGBColor.from_string(DARK)
                    if header and r_idx == 0:
                        run.bold = True


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, side: str, color: str, size="18", space="5") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), size)
    edge.set(qn("w:space"), space)
    edge.set(qn("w:color"), color)
    borders.append(edge)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    run.font.size = Pt(8.5)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    tail = paragraph.add_run(" 页")
    tail.font.size = Pt(8.5)


def add_code_block(doc, text: str) -> None:
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, LIGHT_GREY)
    set_paragraph_border(p, "left", MID_GREY, size="10", space="4")
    for idx, line in enumerate(text.splitlines()):
        if idx:
            p.add_run().add_break()
        p.add_run(line)


def add_callout(doc, title: str, body: str, fill=PALE_TEAL, accent=TEAL) -> None:
    p = doc.add_paragraph(style="Callout")
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, "left", accent)
    run = p.add_run(title + "  ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(accent)
    p.add_run(body)


def add_bullet(doc, text: str, level=0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    doc.add_paragraph(text, style=style)


def new_numbering_id(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(tabs)
    p_pr.append(ind)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(lvl)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered_items(doc: Document, items) -> None:
    num_id = new_numbering_id(doc)
    for text in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.space_after = Pt(4)
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_el])
        p_pr.append(num_pr)
        p.add_run(text)


def add_formula(doc, label: str, formula: str, note: str | None = None) -> None:
    p = doc.add_paragraph(style="Formula")
    p.add_run(label + "\n").bold = True
    r = p.add_run(formula)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    if note:
        p.add_run("\n" + note)


def add_table(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    set_repeat_table_header(table.rows[0])
    format_table(table, widths, font_size=font_size)
    return table


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 28, DARK_BLUE, 0, 12),
        ("Subtitle", 13, MUTED, 0, 12),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    styles["Heading 1"].paragraph_format.page_break_before = False

    for name in ("List Bullet", "List Bullet 2", "List Number", "List Number 2"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
    styles["List Bullet"].paragraph_format.left_indent = Inches(0.375)
    styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.188)
    styles["List Number"].paragraph_format.left_indent = Inches(0.375)
    styles["List Number"].paragraph_format.first_line_indent = Inches(-0.188)

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(DARK)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.05
    set_paragraph_shading(code.element, LIGHT_GREY) if False else None

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    callout.font.name = "Calibri"
    callout._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    callout.font.size = Pt(10)
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.line_spacing = 1.2

    if "Formula" not in styles:
        formula = styles.add_style("Formula", WD_STYLE_TYPE.PARAGRAPH)
    else:
        formula = styles["Formula"]
    formula.font.name = "Calibri"
    formula._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    formula.font.size = Pt(10)
    formula.paragraph_format.left_indent = Inches(0.18)
    formula.paragraph_format.right_indent = Inches(0.18)
    formula.paragraph_format.space_before = Pt(4)
    formula.paragraph_format.space_after = Pt(8)
    formula.paragraph_format.line_spacing = 1.15


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(42)
    set_paragraph_border(p, "bottom", TEAL, size="36", space="12")
    r = p.add_run("QUIETMEL  /  MODEL HANDOFF")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(TEAL)

    title = doc.add_paragraph(style="Title")
    title.add_run("历史人流预测模型\n训练操作指南")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("组员交付版 · hourly-gradient-v1")

    doc.add_paragraph("本指南用于训练一个轻量、可解释的小时级人流预测模型。交付物将由项目负责人导入 Supabase，并通过 Express API 提供给前端。")

    meta = add_table(
        doc,
        ["项目", "职责范围", "数据源", "文档版本"],
        [["QuietMel", "仅模型训练与验证", "City of Melbourne Pedestrian Counting System", "1.0 / 2026-08-06"]],
        [1500, 2100, 3900, 1860],
        font_size=9,
    )
    for cell in meta.rows[0].cells:
        set_cell_shading(cell, DARK_BLUE)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()
    add_callout(
        doc,
        "一句话目标",
        "把近 3 个完整年份的小时人流记录，按“传感器 × 星期 × 季节 × 小时”聚合，输出每组的基准值、区间、梯度和质量标记。",
    )


def build() -> None:
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False
    doc.settings.odd_and_even_pages_header_footer = True

    add_cover(doc)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.even_page_header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.even_page_footer.is_linked_to_previous = False

    header = section.header.paragraphs[0]
    header.text = "QUIETMEL  /  历史人流预测模型训练操作指南"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(MUTED)
    set_paragraph_border(header, "bottom", MID_GREY, size="6", space="6")
    add_page_number(section.footer.paragraphs[0])

    even_header = section.even_page_header.paragraphs[0]
    even_header.text = "QUIETMEL  /  历史人流预测模型训练操作指南"
    even_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in even_header.runs:
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(MUTED)
    set_paragraph_border(even_header, "bottom", MID_GREY, size="6", space="6")
    add_page_number(section.even_page_footer.paragraphs[0])

    # Keep正文页顶端简洁；不同 Word 渲染器对分节后的奇偶页眉处理并不一致。
    for header_para in (header, even_header):
        header_para.clear()
        p_pr = header_para._p.get_or_add_pPr()
        borders = p_pr.find(qn("w:pBdr"))
        if borders is not None:
            p_pr.remove(borders)

    doc.add_heading("1. 任务边界与最终交付", level=1)
    doc.add_paragraph("你负责把历史数据处理为固定格式的模型文件，并完成离线验证。项目负责人负责数据库、上传、后端接口和前端展示。")
    doc.add_heading("你需要完成", level=2)
    for item in (
        "下载并整理官方历史小时人流数据。",
        "按本指南规定的时间、单位和分组规则完成清洗与特征计算。",
        "编写可重复执行的训练脚本，不使用 Notebook 作为唯一交付物。",
        "输出 3 个规定文件，并完成字段、重复值、空值和指标检查。",
    ):
        add_bullet(doc, item)
    doc.add_heading("你不需要完成", level=2)
    for item in (
        "Supabase 建表、权限策略（RLS）或数据导入。",
        "Express 路由、Vercel 部署或前端图表。",
        "实时 API 接入、用户登录或地图功能。",
    ):
        add_bullet(doc, item)
    add_callout(doc, "验收原则", "同一份输入数据和相同参数必须得到一致的输出；所有字段名、枚举值和单位必须严格遵守本指南。", PALE_YELLOW, "9A6A00")

    doc.add_heading("最终交付目录", level=2)
    add_code_block(doc, "model/\n  train_crowd_profiles.py\n  requirements.txt\n  README.md\n  output/\n    crowd_hourly_profiles.csv\n    model_metadata.json\n    validation_metrics.csv")

    doc.add_heading("2. 数据来源与原始字段", level=1)
    doc.add_paragraph("使用 City of Melbourne Open Data 的官方数据集：Pedestrian Counting System — Monthly Counts per Hour。")
    add_callout(doc, "官方数据集", "https://data.melbourne.vic.gov.au/explore/dataset/pedestrian-counting-system-monthly-counts-per-hour/", PALE_BLUE, BLUE)
    doc.add_paragraph("优先下载 Parquet；若官方只提供 CSV，也可使用 CSV。不要使用当前前端的实时接口响应代替历史训练数据。")

    fields = [
        ("id", "忽略", "数据平台记录 ID，不参与建模"),
        ("location_id", "必需", "传感器/地点稳定标识，按字符串读取"),
        ("sensing_date", "必需", "观测日期，解析为本地日期"),
        ("hourday", "必需", "小时，整数 0–23"),
        ("direction_1", "忽略", "方向 1 计数，不用于 v1"),
        ("direction_2", "忽略", "方向 2 计数，不用于 v1"),
        ("pedestriancount", "必需", "该传感器在该小时的总人数"),
        ("sensor_name", "可选", "只用于人工核对，不作为主键"),
        ("location", "可选", "只用于人工核对位置"),
    ]
    add_table(doc, ["原始字段", "状态", "处理规则"], fields, [2100, 1200, 6060], font_size=8.5)

    doc.add_heading("数据时间范围", level=2)
    doc.add_paragraph("每次训练默认选择当前年份之前的 3 个完整年份：")
    add_formula(doc, "选择规则", "start_year = current_year - 3\nend_year   = current_year - 1", "例如在 2026 年执行时，使用 2023-01-01 至 2025-12-31。不要把不完整的当前年份混入训练集。")

    doc.add_heading("3. 固定数据约定", level=1)
    conventions = [
        ("时区", "Australia/Melbourne", "日期、星期和小时均以墨尔本本地时间解释"),
        ("day_of_week", "0–6", "0=Sunday，1=Monday，…，6=Saturday"),
        ("hour_of_day", "0–23", "直接来源于 hourday"),
        ("summer", "12, 1, 2", "澳大利亚季节"),
        ("autumn", "3, 4, 5", "澳大利亚季节"),
        ("winter", "6, 7, 8", "澳大利亚季节"),
        ("spring", "9, 10, 11", "澳大利亚季节"),
        ("流量单位", "persons/minute", "输出统一为每分钟平均人数（ppm）"),
    ]
    add_table(doc, ["项目", "固定值", "说明"], conventions, [1700, 2100, 5560], font_size=8.5)
    add_formula(doc, "小时总数转为分钟平均值", "observed_ppm = pedestriancount / 60", "不要把小时总人数直接写入 baseline_ppm。")

    doc.add_heading("4. 环境准备与运行方式", level=1)
    doc.add_paragraph("建议使用 Python 3.11 或更高版本。训练脚本必须可从命令行独立执行。")
    doc.add_heading("requirements.txt 最低依赖", level=2)
    add_code_block(doc, "pandas>=2.2,<3\nnumpy>=1.26,<3\npyarrow>=15,<22")
    doc.add_heading("命令行接口", level=2)
    add_code_block(doc, "python train_crowd_profiles.py --input pedestrian-counts.parquet --output-dir output")
    doc.add_paragraph("脚本需自动创建 output 目录，并在控制台打印：输入记录数、保留记录数、日期范围、传感器数量、输出分组数、被标记为 low_sample/fallback 的组数。")

    doc.add_heading("5. 数据清洗步骤", level=1)
    add_numbered_items(doc, (
        "读取文件时把 location_id 按字符串处理；只保留必需字段及可选核对字段。",
        "解析 sensing_date；无法解析的记录删除并计数。",
        "把 hourday 转为整数，删除不在 0–23 范围内的记录。",
        "把 pedestriancount 转为数值；删除空值和负数。零值允许保留。",
        "按默认规则筛选最近 3 个完整年份。",
        "使用 sensing_date 生成 day_of_week 和 season；使用 hourday 生成 hour_of_day。",
        "计算 observed_ppm = pedestriancount / 60。",
        "检查重复记录。若同一 location_id、sensing_date、hourday 出现多行，先记录数量，再取 pedestriancount 的总和；不得静默丢弃。",
        "保留清洗统计，用于写入 model_metadata.json。",
    ))
    add_callout(doc, "禁止项", "不得用 sensor_name 代替 location_id；不得把方向计数与 pedestriancount 再次相加；不得对负值取绝对值。", PALE_RED, "B42318")

    doc.add_heading("6. 分组与特征计算", level=1)
    doc.add_paragraph("训练数据的基础分组键固定为：")
    add_code_block(doc, "location_id × day_of_week × season × hour_of_day")
    doc.add_paragraph("每个分组使用其中全部 observed_ppm 样本计算下列特征：")
    features = [
        ("baseline_ppm", "median(observed_ppm)", "稳健基准预测值"),
        ("trimmed_mean_ppm", "mean(remove bottom/top 10%)", "去掉两端各 10% 后的平均值"),
        ("p25_ppm", "25th percentile", "预期低位区间"),
        ("p75_ppm", "75th percentile", "预期高位区间"),
        ("sample_count", "count(observed_ppm)", "该组有效样本数"),
    ]
    add_table(doc, ["输出字段", "计算方法", "用途"], features, [2300, 3200, 3860], font_size=8.5)
    doc.add_paragraph("分位数使用 pandas/numpy 默认线性插值即可，但 README 中必须注明。若样本数小于 10，trimmed mean 仍可计算；切片后为空时退回普通平均值。")

    doc.add_heading("7. 梯度模型", level=1)
    doc.add_paragraph("v1 不训练复杂机器学习模型。每个小时段使用当前基准值和下一小时基准值之间的斜率表示趋势。")
    add_formula(doc, "小时梯度", "gradient_ppm_per_hour = next_hour_baseline_ppm - baseline_ppm")
    add_formula(doc, "小时内预测方程", "predicted_ppm(t) = baseline_ppm + gradient_ppm_per_hour × t", "t = minute / 60，范围为 0 ≤ t < 1。例如 10:30 的 t=0.5。")

    doc.add_heading("下一小时键的计算", level=2)
    for item in (
        "hour 0–22：下一小时仍属于同一天、同一星期、同一季节。",
        "hour 23：下一小时为次日 hour 0；day_of_week 加 1，并按 7 取模。",
        "跨季节边界：用实际日期推导下一小时的 season。不要假设次日季节一定相同。",
        "如果下一小时基准值不可用，gradient_ppm_per_hour 设为 0，并将质量标记提升为 fallback。",
    ):
        add_bullet(doc, item)
    add_callout(doc, "实现提示", "为正确处理跨日和跨季节，建议先建立“日期样本 → 下一小时时间键”的映射，再连接到聚合结果；不要只用 hour+1。", PALE_YELLOW, "9A6A00")

    doc.add_heading("8. 低样本与回退策略", level=1)
    quality_rows = [
        (">= 10", "ok", "直接使用完整分组统计"),
        ("5–9", "low_sample", "仍输出完整分组统计，但标记样本不足"),
        ("< 5", "fallback", "依次使用下面的回退序列"),
    ]
    add_table(doc, ["sample_count", "quality_flag", "处理"], quality_rows, [1800, 1800, 5760], font_size=9)
    doc.add_heading("回退顺序", level=2)
    add_numbered_items(doc, (
        "同一传感器 + 同一 day_of_week + 同一 hour_of_day，忽略 season。",
        "同一传感器 + weekday/weekend + 同一 hour_of_day，忽略具体星期和季节。weekday=Monday–Friday；weekend=Saturday/Sunday。",
        "同一传感器 + 同一 hour_of_day，使用全部可用日期。",
        "仍无样本时保留该键，baseline/区间/梯度写 0，quality_flag=fallback，并在 metadata 中记录数量。",
    ))
    doc.add_paragraph("回退后的数值写入输出字段，但 sample_count 仍保留原始完整分组的样本数，以便下游知道为何发生回退。")

    doc.add_heading("9. 主输出文件：crowd_hourly_profiles.csv", level=1)
    doc.add_paragraph("文件编码必须为 UTF-8；列顺序、名称和枚举值必须完全一致。")
    add_code_block(doc, "location_id,day_of_week,season,hour_of_day,baseline_ppm,trimmed_mean_ppm,p25_ppm,p75_ppm,gradient_ppm_per_hour,sample_count,quality_flag,source_start_date,source_end_date,model_version")

    schema_rows = [
        ("location_id", "text", "非空；来自原始 location_id"),
        ("day_of_week", "integer", "0–6；0=Sunday"),
        ("season", "text", "summer/autumn/winter/spring"),
        ("hour_of_day", "integer", "0–23"),
        ("baseline_ppm", "decimal", "非负；最多 4 位小数"),
        ("trimmed_mean_ppm", "decimal", "非负；最多 4 位小数"),
        ("p25_ppm", "decimal", "非负；最多 4 位小数"),
        ("p75_ppm", "decimal", "非负；最多 4 位小数"),
        ("gradient_ppm_per_hour", "decimal", "可正可负；最多 4 位小数"),
        ("sample_count", "integer", "原始完整分组有效样本数"),
        ("quality_flag", "text", "ok/low_sample/fallback"),
        ("source_start_date", "date", "YYYY-MM-DD"),
        ("source_end_date", "date", "YYYY-MM-DD"),
        ("model_version", "text", "固定 hourly-gradient-v1"),
    ]
    add_table(doc, ["字段", "类型", "规则"], schema_rows, [2450, 1500, 5410], font_size=8)

    doc.add_heading("唯一键", level=2)
    add_code_block(doc, "location_id + day_of_week + season + hour_of_day + model_version")
    doc.add_paragraph("该组合不得重复。数值字段不得出现 NaN、Infinity、-Infinity 或科学计数法。")

    doc.add_heading("示例记录", level=2)
    add_code_block(doc, "37,1,summer,8,12.7500,13.0241,9.5000,16.1000,2.3000,38,ok,2023-01-01,2025-12-31,hourly-gradient-v1")

    doc.add_heading("10. 元数据文件：model_metadata.json", level=1)
    doc.add_paragraph("该文件说明训练范围、规则和清洗结果。字段可扩展，但以下字段必须存在。")
    add_code_block(doc, '{\n  "model_version": "hourly-gradient-v1",\n  "generated_at": "2026-08-06T10:00:00+10:00",\n  "timezone": "Australia/Melbourne",\n  "source_dataset": "pedestrian-counting-system-monthly-counts-per-hour",\n  "source_url": "https://data.melbourne.vic.gov.au/explore/dataset/pedestrian-counting-system-monthly-counts-per-hour/",\n  "source_start_date": "2023-01-01",\n  "source_end_date": "2025-12-31",\n  "input_row_count": 0,\n  "clean_row_count": 0,\n  "dropped_row_count": 0,\n  "duplicate_group_count": 0,\n  "sensor_count": 0,\n  "profile_row_count": 0,\n  "quality_counts": {"ok": 0, "low_sample": 0, "fallback": 0},\n  "unit": "persons_per_minute",\n  "day_of_week_convention": "0=Sunday,...,6=Saturday",\n  "season_convention": "Australian meteorological seasons",\n  "gradient_method": "next_hour_baseline_minus_current_hour_baseline"\n}')

    doc.add_heading("11. 离线验证", level=1)
    doc.add_paragraph("验证只用于判断模型是否可接受，不改变输出字段。把时间范围内最近 4 周作为测试集，其余记录作为训练集。")
    add_numbered_items(doc, (
        "按 sensing_date 排序，找到最大日期。",
        "测试集 = 最大日期往前 27 天至最大日期；训练集 = 更早的记录。",
        "只用训练集重新计算 profile。",
        "为测试集每条记录按 location_id/day_of_week/season/hour_of_day 查找预测值，并应用相同回退规则。",
        "比较 predicted_ppm 与 observed_ppm，按传感器计算 MAE 和 WAPE。",
    ))
    add_formula(doc, "MAE", "mae_ppm = mean(abs(observed_ppm - predicted_ppm))")
    add_formula(doc, "WAPE", "wape_percent = sum(abs(observed_ppm - predicted_ppm)) / sum(observed_ppm) × 100", "当 observed_ppm 总和为 0 时，WAPE 留空并在质量标记中说明。")

    doc.add_heading("validation_metrics.csv 固定字段", level=2)
    add_code_block(doc, "location_id,mae_ppm,wape_percent,test_sample_count,quality_flag")
    doc.add_paragraph("quality_flag 建议规则：测试样本数 >= 10 为 ok，5–9 为 low_sample，<5 为 fallback。指标最多保留 4 位小数。")
    add_callout(doc, "说明", "v1 不设强制准确率门槛。验收重点是可重复、无数据泄漏、单位正确、回退透明。若个别传感器 WAPE 异常高，应在 README 中列出并说明。", PALE_BLUE, BLUE)

    doc.add_heading("12. 自动检查与交付前验收", level=1)
    checks = [
        ("文件存在", "3 个 output 文件均已生成"),
        ("列契约", "列名和列顺序完全一致"),
        ("唯一性", "主输出唯一键无重复"),
        ("取值范围", "day 0–6、hour 0–23、season 与 quality_flag 枚举合法"),
        ("数值质量", "无 NaN/Infinity；除 gradient 外均非负"),
        ("日期", "YYYY-MM-DD，且为 3 个完整年份"),
        ("精度", "小数最多 4 位"),
        ("可重复", "相同输入连续执行两次，输出内容一致"),
        ("说明文件", "README 包含运行命令、依赖、时间范围和异常说明"),
    ]
    add_table(doc, ["检查项", "通过标准"], checks, [2300, 7060], font_size=9)

    doc.add_heading("建议在脚本末尾执行的断言", level=2)
    add_code_block(doc, "assert profiles[KEY_COLUMNS].duplicated().sum() == 0\nassert profiles['day_of_week'].between(0, 6).all()\nassert profiles['hour_of_day'].between(0, 23).all()\nassert set(profiles['season']) <= {'summer', 'autumn', 'winter', 'spring'}\nassert set(profiles['quality_flag']) <= {'ok', 'low_sample', 'fallback'}\nassert np.isfinite(profiles[NUMERIC_COLUMNS].to_numpy()).all()")

    doc.add_heading("13. README 必须写明的内容", level=1)
    for item in (
        "Python 版本、安装命令和完整运行命令。",
        "输入文件来源、下载日期、输入格式和时间范围。",
        "字段映射、时区、星期编码、季节规则和 ppm 单位。",
        "中位数、10% trimmed mean、分位数和梯度的计算方法。",
        "重复记录、缺失值、低样本和下一小时缺失的处理方式。",
        "验证集划分方式、整体或典型 MAE/WAPE，以及已知异常。",
        "输出目录和 3 个交付文件的用途。",
    ):
        add_bullet(doc, item)

    doc.add_heading("14. 交付清单", level=1)
    checklist = [
        ("□", "train_crowd_profiles.py 可在空环境按 README 运行"),
        ("□", "requirements.txt 只包含实际使用的依赖"),
        ("□", "crowd_hourly_profiles.csv 通过所有自动检查"),
        ("□", "model_metadata.json 数字与实际输出一致"),
        ("□", "validation_metrics.csv 覆盖所有可验证的 location_id"),
        ("□", "README.md 已记录运行方法、规则和异常"),
        ("□", "压缩包中不含原始密钥、Supabase key 或 .env 文件"),
    ]
    add_table(doc, ["状态", "交付要求"], checklist, [900, 8460], font_size=9.5)
    add_callout(doc, "提交给项目负责人", "提交整个 model 文件夹。项目负责人将检查输出契约后创建数据库表、上传数据，并完成 Express 与前端对接。", PALE_TEAL, TEAL)

    doc.add_heading("附录 A：下游如何使用这些字段（无需组员实现）", level=1)
    downstream = [
        ("baseline_ppm", "作为当前小时预测基准"),
        ("gradient_ppm_per_hour", "根据当前分钟修正预测趋势"),
        ("p25_ppm / p75_ppm", "前端显示正常预测区间"),
        ("quality_flag", "提示低样本或回退结果"),
        ("model_version", "支持后续模型版本切换与回滚"),
        ("source_start/end_date", "显示模型数据覆盖范围"),
    ]
    add_table(doc, ["字段", "后续用途"], downstream, [3000, 6360], font_size=9)
    doc.add_paragraph("后端查询时会使用当前传感器、星期、季节和小时定位一条 profile，再按当前分钟计算 predicted_ppm。所有命名和单位之所以固定，是为了避免导入后再次转换。")

    doc.add_heading("附录 B：最终数据流", level=1)
    flow = add_table(
        doc,
        ["1. 官方历史数据", "2. 训练脚本", "3. 固定输出", "4. 项目负责人"],
        [["小时人数记录", "清洗、聚合、梯度、验证", "CSV + JSON + 指标", "Supabase → Express → 前端"]],
        [2100, 2500, 2200, 2560],
        font_size=8.8,
    )
    for cell in flow.rows[0].cells:
        set_cell_shading(cell, TEAL)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = end.add_run("END OF HANDOFF GUIDE")
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
